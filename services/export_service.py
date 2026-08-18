"""
ExportService — genera los artefactos finales en los formatos que el
enunciado indica en su stack tecnológico de referencia: PDD/SDD en Word
(python-docx) y la planilla de Estimación en Excel (openpyxl). QA se exporta
también como Excel por tratarse de una matriz tabular.

Este módulo NO decide el contenido (eso es responsabilidad de las Gems):
solo vuelca el texto ya generado sobre un layout con la identidad visual
mínima de un template corporativo (encabezado, título, cuerpo).
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from services.estimacion_service import EstimacionService, FASES_ESPERADAS
from services.pdd_service import PddService
from services.sdd_service import SddService

# Paleta tomada del template histórico de Estimación (Google Sheets → Excel)
_HEADER_FILL = PatternFill(start_color="C9DAF8", end_color="C9DAF8", fill_type="solid")
_TOTAL_FILL = PatternFill(start_color="073763", end_color="073763", fill_type="solid")
_TOTAL_FONT = Font(bold=True, color="FFFFFF")
_HEADER_FONT = Font(bold=True, color="000000")
_THIN = Side(style="thin", color="D0D0D0")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)

# Mismo celeste corporativo que las tablas del template PDD (SAMAN)
_DOCX_HEADER_SHADE = "C9DAF8"


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _tabla_docx(doc, headers: list, filas: list) -> None:
    """Crea una tabla estilo 'Table Grid' con encabezado celeste, igual que
    las tablas del template PDD corporativo. `filas` es una lista de listas
    de strings, en el mismo orden que `headers`."""
    tabla = doc.add_table(rows=1, cols=len(headers))
    tabla.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tabla.rows[0].cells[i]
        cell.text = h
        _shade_cell(cell, _DOCX_HEADER_SHADE)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for fila in filas:
        row_cells = tabla.add_row().cells
        for i, val in enumerate(fila):
            row_cells[i].text = "" if val in (None, "") else str(val)
    doc.add_paragraph("")


def _docx_desde_texto(titulo: str, subtitulo: str, cuerpo: str) -> io.BytesIO:
    doc = Document()

    h = doc.add_heading(titulo, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x5C)

    p = doc.add_paragraph(subtitulo)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph(f"Generado por AutoDocs AI — {datetime.now():%d/%m/%Y %H:%M}")
    doc.add_paragraph("")

    for bloque in (cuerpo or "").split("\n"):
        linea = bloque.strip()
        if not linea:
            doc.add_paragraph("")
        elif linea.startswith("### "):
            doc.add_heading(linea[4:], level=3)
        elif linea.startswith("## "):
            doc.add_heading(linea[3:], level=2)
        elif linea.startswith("# "):
            doc.add_heading(linea[2:], level=1)
        elif linea.startswith(("- ", "* ")):
            doc.add_paragraph(linea[2:], style="List Bullet")
        else:
            doc.add_paragraph(linea)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _pdd_a_docx_estructurado(proyecto: dict, data: dict) -> io.BytesIO:
    """Arma el PDD replicando las secciones del template corporativo
    (SAMAN-style): Introducción, Descripción AS IS (tabla general +
    aplicaciones + camino feliz), Descripción TO BE (alcance, excepciones
    de negocio, errores de aplicación, reportes)."""
    doc = Document()
    nombre_proceso = data.get("nombre_proceso") or proyecto.get("proceso", "")

    h = doc.add_heading("Process Definition Document", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x5C)
    doc.add_heading(f"{proyecto.get('cliente','')} — {nombre_proceso}", level=1)
    doc.add_paragraph(f"Generado por AutoDocs AI — {datetime.now():%d/%m/%Y %H:%M}").italic = True
    doc.add_paragraph("")

    # ── 1. INTRODUCCIÓN ─────────────────────────────────────────────
    doc.add_heading("1. Introducción", level=1)
    doc.add_heading("1.1 Propósito", level=2)
    doc.add_paragraph(
        "El presente documento de definición de procesos (PDD) describe el proceso de "
        "negocio elegido para la automatización, la secuencia de acciones AS IS y la "
        "propuesta TO BE resultante de la preparación para la automatización."
    )
    doc.add_heading("1.2 Objetivos", level=2)
    doc.add_paragraph(data.get("objetivos") or "—")

    # ── 2. DESCRIPCIÓN DEL PROCESO AS IS ────────────────────────────
    doc.add_heading("2. Descripción del proceso AS IS", level=1)
    doc.add_heading("2.1 Descripción general del proceso", level=2)
    campos_generales = [
        ("Nombre del Proceso", data.get("nombre_proceso")),
        ("Área del Proceso", data.get("area_proceso")),
        ("Área", data.get("area")),
        ("Descripción breve", data.get("descripcion_breve")),
        ("Roles/aplicaciones requeridos", data.get("roles_aplicaciones")),
        ("Horario y frecuencia del proceso", data.get("horario_frecuencia")),
        ("# de veces de ejecución", data.get("veces_ejecucion")),
        ("Tiempo de ejecución del proceso", data.get("tiempo_ejecucion")),
        ("Restricciones del proceso", data.get("restricciones")),
        ("Periodo(s) pico", data.get("periodo_pico")),
        ("Volumen máximo de periodo pico", data.get("volumen_pico")),
        ("# de personas que realizan el proceso", data.get("personas_proceso")),
        ("Descripción de los datos de entrada", data.get("datos_entrada")),
        ("Descripción de los datos de salida", data.get("datos_salida")),
    ]
    _tabla_docx(doc, ["Item", "Descripción / Respuesta"], campos_generales)

    doc.add_heading("2.2 Aplicaciones usadas", level=2)
    apps = data.get("aplicaciones") or []
    _tabla_docx(
        doc,
        ["Nombre de aplicación", "Versión", "Idioma", "Environment/Access", "Comentario"],
        [[a.get("nombre"), a.get("version"), a.get("idioma"), a.get("acceso"), a.get("comentario")]
         for a in apps if isinstance(a, dict)] or [["—", "", "", "", ""]],
    )

    doc.add_heading("2.3 Mapa de procesos a nivel detallado — Camino Feliz", level=2)
    pasos = data.get("camino_feliz") or []
    _tabla_docx(
        doc,
        ["#", "Descripción", "Resultado esperado", "Regla de Negocio", "Comentarios"],
        [[p.get("numero"), p.get("descripcion"), p.get("resultado_esperado"),
          p.get("regla_negocio"), p.get("comentarios")]
         for p in pasos if isinstance(p, dict)] or [["—", "", "", "", ""]],
    )

    # ── 3. DESCRIPCIÓN DEL PROCESO TO BE ────────────────────────────
    doc.add_heading("3. Descripción del proceso TO BE", level=1)

    doc.add_heading("3.1 Dentro del alcance de RPA", level=2)
    dentro = data.get("dentro_alcance") or []
    if dentro:
        for item in dentro:
            doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph("—")

    doc.add_heading("3.2 Fuera del alcance de RPA", level=2)
    fuera = data.get("fuera_alcance") or []
    _tabla_docx(
        doc, ["Actividad / Acción", "Motivo fuera de alcance"],
        [[f.get("actividad"), f.get("motivo")] for f in fuera if isinstance(f, dict)] or [["Ninguna", ""]],
    )

    doc.add_heading("3.3 Manejo de Excepciones", level=2)
    doc.add_heading("3.3.1 Excepciones del Negocio conocidas", level=3)
    exc_neg = data.get("excepciones_negocio_conocidas") or []
    _tabla_docx(
        doc, ["Nombre de la excepción", "Acción", "Parámetros", "Acción a realizar"],
        [[e.get("nombre"), e.get("accion"), e.get("parametros"), e.get("accion_a_realizar")]
         for e in exc_neg if isinstance(e, dict)] or [["—", "", "", ""]],
    )
    doc.add_heading("3.3.2 Excepciones de negocio no conocidas", level=3)
    doc.add_paragraph(data.get("excepciones_negocio_desconocidas") or
                       "Capturar pantalla y notificar por correo; el robot continúa con la siguiente transacción.")

    doc.add_heading("3.4 Manejo de excepciones y errores de aplicaciones", level=2)
    doc.add_heading("3.4.1 Errores y Excepciones Conocidas de las Aplicaciones", level=3)
    err_sis = data.get("errores_sistema_conocidos") or []
    _tabla_docx(
        doc, ["Nombre de Excepción/Error", "Acción", "Parámetros", "Acción a ejecutar"],
        [[e.get("nombre"), e.get("accion"), e.get("parametros"), e.get("accion_a_ejecutar")]
         for e in err_sis if isinstance(e, dict)] or [["—", "", "", ""]],
    )
    doc.add_heading("3.4.2 Errores y Excepciones Desconocidas de las Aplicaciones", level=3)
    doc.add_paragraph(data.get("errores_sistema_desconocidos") or
                       "Reintentar acceso a la aplicación 3 veces y luego finalizar el subproceso.")

    doc.add_heading("3.5 Reportes", level=2)
    reportes = data.get("reportes") or []
    _tabla_docx(
        doc, ["Tipo de reporte", "Frecuencia de actualización", "Detalle"],
        [[r.get("tipo"), r.get("frecuencia"), r.get("detalle")] for r in reportes if isinstance(r, dict)]
        or [["Logs del proceso", "Diario", ""]],
    )

    # ── notas del análisis automatizado (transparencia) ─────────────
    if data.get("supuestos") or data.get("preguntas_abiertas"):
        doc.add_heading("4. Notas del análisis automatizado", level=1)
        if data.get("supuestos"):
            doc.add_heading("4.1 Supuestos", level=2)
            for s in data["supuestos"]:
                doc.add_paragraph(str(s), style="List Bullet")
        if data.get("preguntas_abiertas"):
            doc.add_heading("4.2 Preguntas abiertas para el cliente", level=2)
            for s in data["preguntas_abiertas"]:
                doc.add_paragraph(str(s), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _sdd_a_docx_estructurado(proyecto: dict, data: dict) -> io.BytesIO:
    """Arma el SDD replicando las secciones y tablas del template técnico
    estándar (estilo UNACEM/SAMAN): historial de revisiones, detalles del
    proceso/ejecución/desarrollo, packages, archivos de flujo, seguimiento,
    mejoras futuras, recursos externos y glosario."""
    doc = Document()
    nombre_robot = data.get("nombre_robot") or proyecto.get("proceso", "")
    rev = data.get("revision") or {}

    h = doc.add_heading("Solution Design Document", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x5C)
    doc.add_heading(f"{proyecto.get('cliente','')} — {nombre_robot}", level=1)
    doc.add_paragraph(f"Generado por AutoDocs AI — {datetime.now():%d/%m/%Y %H:%M}").italic = True
    doc.add_paragraph("")

    doc.add_heading("Historial de revisiones del documento", level=2)
    _tabla_docx(
        doc, ["Fecha", "Versión", "Descripción", "Autor", "Revisor", "Aprobador", "Cargo"],
        [[rev.get("fecha") or datetime.now().strftime("%m/%Y"), rev.get("version", "1.0"),
          rev.get("descripcion", "Primera versión del documento"), rev.get("autor"),
          rev.get("revisor"), rev.get("aprobador"), rev.get("cargo")]],
    )

    # ── 1-3. Introducción / Objetivos / Contactos ───────────────────
    doc.add_heading("1. Introducción", level=2)
    doc.add_paragraph(data.get("introduccion") or
                       "En el presente documento se da a conocer el diseño de la automatización "
                       "RPA aplicada para el proceso. El mismo debe ser actualizado con cada "
                       "modificación que se realice a futuro.")
    doc.add_heading("2. Objetivos", level=2)
    doc.add_paragraph(data.get("objetivos") or
                       "Informar al implementador o futuro soporte de la plataforma el proceso automático.")
    doc.add_heading("3. Email de Contactos", level=2)
    doc.add_paragraph(f"Analista del proceso: {data.get('contacto_analista','') } — {data.get('contacto_analista_mail','')}")
    doc.add_paragraph(f"Desarrolladores RPA: {data.get('contacto_dev','')} — {data.get('contacto_dev_mail','')}")
    doc.add_paragraph(f"Referente del proceso (Cliente): {data.get('contacto_cliente','')} — {data.get('contacto_cliente_mail','')}")

    # ── 4. Detalles del Proceso ──────────────────────────────────────
    doc.add_heading("4. Detalles del Proceso", level=2)
    _tabla_docx(
        doc, ["Item", "Descripción"],
        [
            ["Nombre del proceso", nombre_robot],
            ["Tipo de Robot", data.get("tipo_robot")],
            ["¿Se usa orquestador?", data.get("usa_orquestador")],
            ["Escalable", data.get("escalable")],
            ["Versión de la plataforma", data.get("version_plataforma")],
        ],
    )

    # ── 5. Detalles de ejecución del proceso ─────────────────────────
    doc.add_heading("5. Detalles de ejecución del proceso", level=2)
    _tabla_docx(
        doc, ["Item", "Descripción"],
        [
            ["Detalle del ambiente de producción", data.get("ambiente_produccion")],
            ["Pre-requisitos para ejecutar", data.get("prerequisitos_ejecutar")],
            ["Datos de entrada", data.get("datos_entrada")],
            ["Datos de salida esperados", data.get("datos_salida")],
            ["¿Cómo empezar el proceso automático?", data.get("como_empezar")],
            ["Reportes", data.get("reportes")],
            ["¿Cómo es usado el orquestador?", data.get("uso_orquestador")],
            ["Política de contraseñas", data.get("politica_contrasenas")],
            ["Credenciales guardadas", data.get("credenciales_guardadas")],
            ["Lista de queues", data.get("lista_queues")],
            ["Detalles del calendario", data.get("detalles_calendario")],
            ["¿Múltiples resoluciones soportadas?", data.get("multiples_resoluciones")],
            ["Resolución recomendada", data.get("resolucion_recomendada")],
        ],
    )

    # ── 6. Detalles del desarrollo ────────────────────────────────────
    doc.add_heading("6. Detalles del desarrollo", level=2)
    _tabla_docx(
        doc, ["Item", "Descripción"],
        [
            ["Ambiente usado para el desarrollo", data.get("ambiente_desarrollo")],
            ["Pre-requisito del ambiente", data.get("prerequisito_ambiente")],
            ["Repositorios", data.get("repositorios")],
            ["Método de configuración", data.get("metodo_configuracion")],
            ["Lista de componentes reutilizados", data.get("componentes_reutilizados")],
            ["Lista de nuevos componentes reutilizables", data.get("nuevos_componentes")],
        ],
    )

    # ── 7. Packages ────────────────────────────────────────────────────
    doc.add_heading("7. Packages", level=2)
    packages = data.get("packages") or []
    _tabla_docx(
        doc, ["Nombre", "Descripción"],
        [[pkg.get("nombre"), pkg.get("descripcion")] for pkg in packages if isinstance(pkg, dict)]
        or [["—", ""]],
    )

    # ── 8. Archivos de Flujo ───────────────────────────────────────────
    doc.add_heading("8. Archivos de Flujo", level=3)
    archivos = data.get("archivos_flujo") or []
    _tabla_docx(
        doc, ["Archivo", "Descripción", "Argumentos"],
        [[a.get("archivo"), a.get("descripcion"), a.get("argumentos")] for a in archivos if isinstance(a, dict)]
        or [["—", "", ""]],
    )

    # ── 9. Diagramas ────────────────────────────────────────────────────
    doc.add_heading("9. Diagramas", level=3)
    doc.add_paragraph("Level 1")
    doc.add_paragraph("Level 2")
    doc.add_paragraph("Level 3")
    p_diag = doc.add_paragraph()
    p_diag.add_run(
        "⚠ Insertar acá los diagramas C4 (Nivel 1/2/3) del flujo de automatización — "
        "AutoDocs AI no genera diagramas de arquitectura de bajo nivel automáticamente."
    ).italic = True

    # ── 10. Seguimiento ──────────────────────────────────────────────
    doc.add_heading("10. Seguimiento", level=2)
    doc.add_heading("10.1 Ejecuciones", level=3)
    doc.add_paragraph(data.get("seguimiento_ejecuciones") or "A definir.")
    doc.add_heading("10.2 Tips de soporte", level=3)
    doc.add_paragraph(data.get("tips_soporte") or "<<Detallar los posibles resultados y las acciones a realizar en cada caso>>")

    # ── 11. Mejoras Futuras ────────────────────────────────────────────
    doc.add_heading("11. Mejoras Futuras", level=2)
    mejoras = data.get("mejoras_futuras") or []
    if mejoras:
        for m in mejoras:
            doc.add_paragraph(str(m), style="List Bullet")
    else:
        doc.add_paragraph("<<Se detallan aquí las mejoras a futuro detectadas durante el desarrollo>>")

    # ── 12. Recursos externos ───────────────────────────────────────────
    doc.add_heading("12. Recursos externos", level=2)
    recursos = data.get("recursos_externos") or []
    _tabla_docx(
        doc, ["Recurso", "Web", "Descripción"],
        [[r.get("recurso"), r.get("web"), r.get("descripcion")] for r in recursos if isinstance(r, dict)]
        or [["—", "", ""]],
    )

    # ── 13. Glosario ─────────────────────────────────────────────────
    doc.add_heading("13. Glosario", level=2)
    glosario = data.get("glosario") or []
    _tabla_docx(
        doc, ["Término", "Descripción"],
        [[g.get("termino"), g.get("descripcion")] for g in glosario if isinstance(g, dict)]
        or [["—", ""]],
    )

    if data.get("supuestos") or data.get("preguntas_abiertas"):
        doc.add_heading("14. Notas del análisis automatizado", level=2)
        if data.get("supuestos"):
            doc.add_heading("14.1 Supuestos", level=3)
            for s in data["supuestos"]:
                doc.add_paragraph(str(s), style="List Bullet")
        if data.get("preguntas_abiertas"):
            doc.add_heading("14.2 Preguntas abiertas para el cliente/implementador", level=3)
            for s in data["preguntas_abiertas"]:
                doc.add_paragraph(str(s), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _estimacion_a_xlsx_estructurada(proyecto: dict, data: dict) -> io.BytesIO:
    """Arma la planilla de Estimación con el mismo layout que el template
    histórico: bloque de fases (col C..J), fila de Horas Total, bloque de
    detalle de Desarrollo (col A..H) y hoja Param con los % base por fase."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Hoja 1"

    headers = ["Fases", "Story", "Task", "Desglose de tareas", "Entregable", "Responsable", "Total Horas", "Días hábiles"]
    for i, h in enumerate(headers):
        c = ws.cell(row=1, column=3 + i, value=h)
        c.fill = _HEADER_FILL
        c.font = _HEADER_FONT

    nombre_proceso = data.get("proceso") or f"{proyecto.get('cliente','')} - {proyecto.get('proceso','')}"
    story = f"Proceso: {nombre_proceso}"

    fases_por_nombre = {f.get("fase"): f for f in data.get("fases", []) if isinstance(f, dict)}
    fila = 2
    for nombre_fase in FASES_ESPERADAS:
        f = fases_por_nombre.get(nombre_fase, {})
        horas = EstimacionService._num(f.get("horas"))
        dias = round(horas / 6, 6) if horas else 0
        ws.cell(row=fila, column=3, value=nombre_fase)
        ws.cell(row=fila, column=4, value=story)
        ws.cell(row=fila, column=5, value=f.get("task", nombre_fase))
        ws.cell(row=fila, column=6, value=f.get("desglose", ""))
        ws.cell(row=fila, column=7, value=f.get("entregable", ""))
        ws.cell(row=fila, column=8, value=f.get("responsable", ""))
        ws.cell(row=fila, column=9, value=horas)
        ws.cell(row=fila, column=10, value=dias)
        if nombre_fase == "Desarrollo":
            ws.cell(row=1, column=1, value=story)
        fila += 1

    total_horas = EstimacionService.horas_totales(data)
    total_dias = round(total_horas / 6, 6) if total_horas else 0

    fila_total = fila  # fila 9 en el layout original
    ws.cell(row=fila_total, column=8, value="Horas Total")
    for col, val in [(9, total_horas), (10, total_dias)]:
        c = ws.cell(row=fila_total, column=col, value=val)
        c.fill = _TOTAL_FILL
        c.font = _TOTAL_FONT

    fila_otras = fila_total + 1
    ws.cell(row=fila_otras, column=1, value="Otras consideraciones")
    if data.get("otras_consideraciones"):
        ws.cell(row=fila_otras + 1, column=1, value=str(data["otras_consideraciones"])[:2000])

    fila_total_proceso = fila_otras + 2
    ws.cell(row=fila_total_proceso, column=1, value=f"TOTAL {story}")
    c = ws.cell(row=fila_total_proceso, column=9, value=total_horas)
    c.fill = _TOTAL_FILL
    c.font = _TOTAL_FONT

    # ── supuestos / preguntas abiertas (transparencia de incertidumbre) ──
    fila_notas = fila_total_proceso + 2
    if data.get("supuestos"):
        ws.cell(row=fila_notas, column=1, value="Supuestos").font = Font(bold=True)
        fila_notas += 1
        for s in data["supuestos"]:
            ws.cell(row=fila_notas, column=1, value=f"- {s}")
            fila_notas += 1
    if data.get("preguntas_abiertas"):
        fila_notas += 1
        ws.cell(row=fila_notas, column=1, value="Preguntas abiertas para el cliente").font = Font(bold=True)
        fila_notas += 1
        for s in data["preguntas_abiertas"]:
            ws.cell(row=fila_notas, column=1, value=f"- {s}")
            fila_notas += 1

    # ── bloque de detalle de Desarrollo (igual al template histórico) ──
    fila_detalle_header = max(fila_notas + 2, 15)
    detalle_headers = ["Proceso", "", "Observación", "Tarea/Proceso", "Complejidad", "Sistemas externos", "Horas Estimadas", "Reutilizando"]
    for i, h in enumerate(detalle_headers):
        if not h:
            continue
        c = ws.cell(row=fila_detalle_header, column=1 + i, value=h)
        c.fill = _HEADER_FILL
        c.font = _HEADER_FONT

    fila_d = fila_detalle_header + 1
    suma_horas_detalle = 0.0
    suma_reutilizando = 0.0
    for t in data.get("detalle_desarrollo", []):
        if not isinstance(t, dict):
            continue
        horas_t = EstimacionService._num(t.get("horas"))
        reutil_t = EstimacionService._num(t.get("reutilizando")) if t.get("reutilizando") not in (None, "") else horas_t
        ws.cell(row=fila_d, column=3, value=t.get("observacion", ""))
        ws.cell(row=fila_d, column=4, value=t.get("tarea", ""))
        ws.cell(row=fila_d, column=5, value=t.get("complejidad", ""))
        ws.cell(row=fila_d, column=6, value=t.get("sistemas_externos", 1))
        ws.cell(row=fila_d, column=7, value=horas_t)
        ws.cell(row=fila_d, column=8, value=reutil_t)
        suma_horas_detalle += horas_t
        suma_reutilizando += reutil_t
        fila_d += 1

    ws.cell(row=fila_d, column=4, value="Total Desarrollo").font = Font(bold=True)
    ws.cell(row=fila_d, column=7, value=suma_horas_detalle).font = Font(bold=True)
    ws.cell(row=fila_d, column=8, value=suma_reutilizando).font = Font(bold=True)

    ws.column_dimensions["A"].width = 18.75
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 17.5
    ws.column_dimensions["D"].width = 45
    ws.column_dimensions["E"].width = 30
    ws.column_dimensions["F"].width = 18
    ws.column_dimensions["G"].width = 18
    ws.column_dimensions["H"].width = 18
    ws.column_dimensions["I"].width = 12
    ws.column_dimensions["J"].width = 12
    ws.freeze_panes = "A2"

    # ── hoja Param: % base por fase, igual al template histórico ──
    wp = wb.create_sheet("Param")
    wp.append(["Fases", "% Tarea Desarrollo"])
    for c in wp[1]:
        c.fill = _HEADER_FILL
        c.font = _HEADER_FONT
    porcentajes = [
        ("Relevamiento", 0.10), ("Análisis y Diseño", 0.10), ("Desarrollo", "-"),
        ("Pruebas", 0.20), ("Ajustes", 0.20), ("Documentación", 0.15), ("Gestión", 0.10),
    ]
    for nombre, pct in porcentajes:
        wp.append([nombre, pct])
    wp.column_dimensions["A"].width = 20
    wp.column_dimensions["B"].width = 18

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _xlsx_desde_texto(titulo: str, cuerpo: str) -> io.BytesIO:
    wb = Workbook()
    ws = wb.active
    ws.title = titulo[:31]

    header_fill = PatternFill(start_color="1A2A5C", end_color="1A2A5C", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    fila = 1
    ws.cell(row=fila, column=1, value=titulo).font = Font(bold=True, size=14)
    fila += 2

    for linea in (cuerpo or "").split("\n"):
        linea = linea.strip()
        if not linea:
            continue
        if "|" in linea:
            celdas = [c.strip(" -") for c in linea.split("|") if c.strip(" -")]
            if not celdas:
                continue
            for col, valor in enumerate(celdas, start=1):
                cell = ws.cell(row=fila, column=col, value=valor)
                if fila == 3 or linea.lower().startswith(("condición", "campo", "excepción")):
                    cell.fill = header_fill
                    cell.font = header_font
        else:
            ws.cell(row=fila, column=1, value=linea.lstrip("-* "))
        fila += 1

    for col in range(1, 8):
        ws.column_dimensions[get_column_letter(col)].width = 28
    ws.freeze_panes = "A4"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


class ExportService:

    @staticmethod
    def pdd_a_docx(proyecto: dict, contenido: str) -> io.BytesIO:
        """La Gem PDD responde en JSON (ver system_instruction en
        config_manager). Se parsea con PddService y se arma el documento
        con las secciones del template corporativo. Si el contenido no es
        interpretable como PDD estructurado, degrada al volcado de texto
        genérico en vez de romper la descarga."""
        data = PddService.parsear(contenido)
        tiene_contenido = bool(
            data.get("nombre_proceso") or data.get("camino_feliz") or data.get("aplicaciones")
        )
        if not tiene_contenido:
            return _docx_desde_texto(
                f"PDD — {proyecto['proceso']}",
                f"Cliente: {proyecto['cliente']} · Tecnología objetivo: {proyecto.get('tecnologia') or '—'}",
                contenido,
            )
        return _pdd_a_docx_estructurado(proyecto, data)

    @staticmethod
    def sdd_a_docx(proyecto: dict, contenido: str) -> io.BytesIO:
        """La Gem SDD responde en JSON (ver system_instruction en
        config_manager). Se parsea con SddService y se arma el documento
        con las secciones y tablas del template técnico estándar. Si el
        contenido no es interpretable como SDD estructurado, degrada al
        volcado de texto genérico en vez de romper la descarga."""
        data = SddService.parsear(contenido)
        tiene_contenido = bool(
            data.get("nombre_robot") or data.get("archivos_flujo") or data.get("packages")
        )
        if not tiene_contenido:
            return _docx_desde_texto(
                f"SDD — {proyecto['proceso']}",
                f"Cliente: {proyecto['cliente']} · Criticidad: {proyecto.get('criticidad') or '—'}",
                contenido,
            )
        return _sdd_a_docx_estructurado(proyecto, data)

    @staticmethod
    def qa_a_xlsx(proyecto: dict, contenido: str) -> io.BytesIO:
        return _xlsx_desde_texto(f"QA — {proyecto['proceso']}", contenido)

    @staticmethod
    def estimacion_a_xlsx(proyecto: dict, contenido: str) -> io.BytesIO:
        """La Gem Estimación responde en JSON (ver system_instruction en
        config_manager). Se parsea con EstimacionService y se arma la
        planilla con el layout del template histórico. Si por algún motivo
        el contenido no es interpretable como la estimación esperada,
        degrada al volcado de texto genérico en vez de romper la descarga."""
        data = EstimacionService.parsear(contenido)
        if EstimacionService.horas_totales(data) == 0 and not data.get("detalle_desarrollo"):
            return _xlsx_desde_texto(f"Estimación — {proyecto['proceso']}", contenido)
        return _estimacion_a_xlsx_estructurada(proyecto, data)
